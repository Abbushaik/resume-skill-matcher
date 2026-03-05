"""
DOCX Resume Parser using python-docx.
Extracts raw text from paragraphs and tables inside .docx files.
"""

import logging
from docx import Document

logger = logging.getLogger(__name__)


def extract_text_from_docx(file_path: str) -> str:
    """
    Extract all text from a DOCX file (paragraphs + table cells).

    Args:
        file_path: Path to the DOCX file.

    Returns:
        Extracted text as a single string.

    Raises:
        FileNotFoundError: If the file does not exist.
        ValueError: If no text could be extracted.
    """
    try:
        doc = Document(file_path)
        lines = []

        # Extract paragraph text
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                lines.append(text)

        # Extract table cell text
        # (many resumes use tables for layout)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text and text not in lines:
                        lines.append(text)

        if not lines:
            raise ValueError("No extractable text found in DOCX file.")

        combined = "\n".join(lines)
        logger.info(f"DOCX parsed: {len(combined)} characters extracted.")
        return combined

    except FileNotFoundError:
        logger.error(f"DOCX file not found: {file_path}")
        raise
    except Exception as e:
        logger.error(f"Failed to parse DOCX '{file_path}': {e}")
        raise